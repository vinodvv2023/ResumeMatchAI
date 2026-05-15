"""
OCR Block Parser
================
Accepts PDF / DOCX / RTF / TXT files and returns structured sections.

Pipeline:
  1. Detect file format → choose extraction strategy
  2. For PDFs: try digital text first (pdfplumber), fallback to OCR
  3. For DOCX: python-docx
  4. For RTF:  striprtf
  5. For TXT:  raw read
  6. Section header detection via regex → group lines into blocks
  7. Return { contact, summary, skills, experience, education,
              certifications, projects, raw_text }
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Any

from backend.config import VISION_MODEL

logger = logging.getLogger(__name__)

# ── Section Header Patterns ──────────────────────────────────────────────────

SECTION_PATTERNS: dict[str, list[str]] = {
    "contact": [
        r"contact(\s+info(rmation)?)?", r"personal(\s+details?)?",
        r"profile", r"about(\s+me)?",
    ],
    "summary": [
        r"(professional\s+)?summary", r"objective", r"overview",
        r"career\s+(summary|objective)", r"professional\s+profile",
    ],
    "skills": [
        r"(technical\s+)?skills?", r"core\s+competenc(y|ies)",
        r"technologies", r"tech\s+stack", r"expertise", r"proficienc(y|ies)",
        r"tools?\s+(&|and)\s+technologies",
    ],
    "experience": [
        r"(work\s+|professional\s+)?experience", r"employment(\s+history)?",
        r"work\s+history", r"career\s+history", r"professional\s+background",
    ],
    "education": [
        r"education(\s+&\s+training)?", r"academic(\s+background)?",
        r"qualifications?", r"degrees?", r"university", r"college",
    ],
    "certifications": [
        r"certifications?", r"certificates?", r"licen[cs]es?",
        r"accreditations?", r"credentials?",
    ],
    "projects": [
        r"projects?", r"portfolio", r"side\s+projects?",
        r"personal\s+projects?", r"open(\s+|-)?source",
    ],
    "languages": [
        r"languages?", r"spoken\s+languages?",
    ],
    "interests": [
        r"interests?", r"hobbies", r"activities",
    ],
    "references": [
        r"references?", r"referees?",
    ],
}

_SECTION_RE = {
    section: re.compile(
        r"^\s*(" + "|".join(pats) + r")\s*[:\-–—]?\s*$",
        re.IGNORECASE,
    )
    for section, pats in SECTION_PATTERNS.items()
}

_HEADER_RE = re.compile(
    r"^\s*([A-Z][A-Z\s&/\-]{2,40})\s*[:\-–—]?\s*$"
)


# ── Format Detection ─────────────────────────────────────────────────────────

def parse_file(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Entry point: detect format and return structured blocks."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    elif ext == ".rtf":
        return _parse_rtf(file_bytes)
    elif ext == ".txt":
        return _parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── PDF Parser ───────────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> dict[str, Any]:
    print(f"[OCR] _parse_pdf received {len(file_bytes)} bytes, first 20 bytes hex: {file_bytes[:20].hex()}")

    import fitz

    links = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            text = page.get_text() or ""
            pages_text.append(text)
            for annot in page.annots():
                if annot.info.get("uri"):
                    links.append(annot.info["uri"])
        doc.close()

        text = "\n".join(pages_text).strip()
        print(f"[OCR] PyMuPDF extracted {len(text)} chars")

        # If very little text was extracted, fall back to OCR
        if len(text) < 100:
            print(f"[OCR] Digital text too short ({len(text)} chars), falling back to OCR")
            text = _pdf_ocr(file_bytes)
            print(f"[OCR] OCR result length: {len(text)} chars")
    except Exception as exc:
        print(f"[OCR] PyMuPDF failed ({exc}), trying OCR + Vision fallback")
        text = _pdf_ocr(file_bytes)
        if len(text) < 100:
            text = _vision_ocr_fallback(file_bytes)

    return _build_blocks(text, links)


def _pdf_ocr(file_bytes: bytes) -> str:
    """Convert PDF pages to images and run Tesseract OCR on each."""
    try:
        import fitz
        import pytesseract
        import numpy as np
        from PIL import Image

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        print(f"[OCR] PyMuPDF OCR: {len(doc)} pages")
        parts: list[str] = []

        for page in doc:
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img, config="--psm 6")
            parts.append(page_text)
        doc.close()

        return "\n".join(parts)
    except ImportError as e:
        print(f"[OCR] ERROR dependencies missing: {e}")
        return ""
    except Exception as exc:
        print(f"[OCR] ERROR failed: {exc}")
        return ""


# ── DOCX Parser ──────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes) -> dict[str, Any]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        lines: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        # Also pull from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())

        text = "\n".join(lines)
        return _build_blocks(text)
    except Exception as exc:
        logger.error("DOCX parsing failed: %s", exc)
        return _empty_blocks("")


# ── RTF Parser ───────────────────────────────────────────────────────────────

def _parse_rtf(file_bytes: bytes) -> dict[str, Any]:
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(file_bytes.decode("latin-1", errors="replace"))
        return _build_blocks(text)
    except Exception as exc:
        logger.error("RTF parsing failed: %s", exc)
        return _empty_blocks("")


# ── TXT Parser ───────────────────────────────────────────────────────────────

def _parse_txt(file_bytes: bytes) -> dict[str, Any]:
    text = file_bytes.decode("utf-8", errors="replace")
    return _build_blocks(text)


# ── Block Detector ───────────────────────────────────────────────────────────

def _build_blocks(raw_text: str, links: list[str] = None) -> dict[str, Any]:
    """
    Detect section headers in raw_text and group lines under each section.
    Returns a structured dict with known sections + raw_text.
    """
    blocks = _empty_blocks(raw_text)
    if links:
        blocks["links"] = links
    
    current_section = "summary"   # assume text before first header = summary
    current_lines: list[str] = []

    for line in raw_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        detected = _detect_section(stripped)
        if detected:
            # Save accumulated lines to the previous section
            if current_lines:
                blocks[current_section].extend(current_lines)
            current_section = detected
            current_lines = []
        else:
            current_lines.append(stripped)

    # Flush remaining lines
    if current_lines:
        blocks[current_section].extend(current_lines)

    # Extract contact info from the top ~10 lines via regex
    blocks["contact"] = _extract_contact(raw_text, blocks["contact"])

    return blocks


def _empty_blocks(raw_text: str) -> dict[str, Any]:
    return {
        "raw_text": raw_text,
        "contact": [],
        "links": [],
        "summary": [],
        "skills": [],
        "experience": [],
        "education": [],
        "certifications": [],
        "projects": [],
        "languages": [],
        "interests": [],
        "references": [],
        "other": [],
    }


def _detect_section(line: str) -> str | None:
    """Return section key if line matches a known section header, else None."""
    for section, pattern in _SECTION_RE.items():
        if pattern.match(line):
            return section
    # Fallback: ALL CAPS line ≥ 3 chars that looks like a header
    if _HEADER_RE.match(line) and line.isupper() and len(line) >= 4:
        for section, pats in SECTION_PATTERNS.items():
            for pat in pats:
                if re.search(pat, line, re.IGNORECASE):
                    return section
    return None


def _extract_contact(raw_text: str, existing: list[str]) -> list[str]:
    """Extract email, phone, LinkedIn, name from top of document."""
    contact = list(existing)
    top = "\n".join(raw_text.splitlines()[:20])

    email_re = re.compile(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", re.IGNORECASE)
    phone_re = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
    linkedin_re = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)

    for m in email_re.findall(top):
        if m not in contact:
            contact.append(f"email: {m}")
    for m in phone_re.findall(top):
        clean = m.strip()
        if len(clean) >= 7 and f"phone: {clean}" not in contact:
            contact.append(f"phone: {clean}")
    for m in linkedin_re.findall(top):
        if m not in contact:
            contact.append(f"linkedin: {m}")

    return contact


def _vision_ocr_fallback(file_bytes: bytes) -> str:
    """Use DeepInfra multimodal model when Tesseract fails."""
    try:
        from backend.services.vision_ocr import extract_text_from_pdf_bytes
        return extract_text_from_pdf_bytes(file_bytes, "uploaded.pdf")
    except Exception as exc:
        print(f"[OCR] Vision fallback failed: {type(exc).__name__}: {exc}")
        return ""
